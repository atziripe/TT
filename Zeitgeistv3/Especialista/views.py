from django.shortcuts import render,  redirect
from django.http import HttpResponse, JsonResponse
from django.conf import settings
from Usuario import views
from .forms import FormEditarE
from Usuario.models import Paciente, Cuidador, Especialista, Administrador
from django.contrib.auth.models import User
from Paciente.models import Ap_Screening, Screening
from Usuario.models import Paciente, Especialista, User
import random, re, json, jwt, requests
import string, datetime
# Para reporte
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx2pdf import convert
import pythoncom 
import pypandoc
from pypandoc.pandoc_download import download_pandoc
import sys, os
import comtypes.client

def inicioEsp(request):
    try:
        return render(request, "Especialista/inicioEspecialista.html")
    except:
        print("No se accedió a la página con credenciales de usuario válidas")
        return render(request, "Usuarios/index.html")

def editE(request, token, tipo, name):
    try:  
        base = "Especialista/baseEspecialista.html" #Para la base de edicion necesitamos tener el menu del perfil que estamos editando
        decodedToken = jwt.decode(token, key=settings.SECRET_KEY, algorithms=['HS256'])
        iduser = decodedToken['user_id']
        print("iduser", iduser)
        infoU = requests.get('http://127.0.0.1:8000/v1/userd/'+str(iduser)+'')
        infoE = requests.get('http://127.0.0.1:8000/v1/Especialistauser/'+str(iduser)+'')

        if infoE.ok and infoU.ok:
            initial_dict = {
                "nvo_nombre":json.loads(infoU.content)['first_name'],
                "nvo_apellidos": json.loads(infoU.content)['last_name'],
                "nvo_nombreUsuario":json.loads(infoU.content)['username'],
                "nvo_correo":json.loads(infoU.content)['email'],
                "nvos_datos_generales":json.loads(infoE.content)['datos_generales'],
                "nvo_numPacientes":json.loads(infoE.content)['numPacientes_Max'], 
            }
        else:
            print("Ocurrio error en usuario", infoU.status_code)
            print("Ocurrio error en especialista", infoE.status_code)
        if request.method=="POST": 
            feditE = FormEditarE(request.POST, initial=initial_dict)
            #try:      
            if feditE.is_valid(): 
                payload = {
                    "username": feditE.cleaned_data['nvo_nombreUsuario'],
                    "first_name": feditE.cleaned_data['nvo_nombre'],
                    "last_name": feditE.cleaned_data['nvo_apellidos'],
                    "email": feditE.cleaned_data['nvo_correo']
                }
                updateU =  requests.put('http://127.0.0.1:8000/v1/editarperfil/'+str(iduser)+'', data=json.dumps(
                            payload), headers={'content-type': 'application/json', "Authorization": "Bearer "+ token +""})
                if updateU.ok:
                    print("Se pudo actualizar el usuario")
                    payloadE = {
                        "datos_generales":feditE.cleaned_data['nvos_datos_generales'],
                        "numPacientes_Max":feditE.cleaned_data['nvo_numPacientes'],
                    }
                    print(payloadE)
                    updateE =requests.put('http://127.0.0.1:8000/v1/editarespecialista/'+str(json.loads(infoE.content)['id']) +'', data=json.dumps(payloadE), headers={'content-type': 'application/json'})
                    if updateE.ok:
                        return render(request, "Especialista/inicioEspecialista.html", {"name":feditE.cleaned_data['nvo_nombre'],"name":feditE.cleaned_data['nvo_nombre'], "tipo": tipo, "access": token, "modified" : True})
                    else:
                        print(updateE.json())
                else:
                    print(updateU.json())
                    print("No se pudo hacer el registro del usuario")
                    return render(request, "Especialista/editarEspecialista.html", {"name":feditE.cleaned_data['nvo_nombre'],"form": feditE, "user": iduser, "access": token, "tipo": tipo, "already_exists": True, "base": base})


        else:
            feditE=FormEditarE(initial=initial_dict)
        return render(request, "Especialista/editarEspecialista.html", {"name": name, "form": feditE, "user": iduser, "base": base, "tipo": tipo, "access": token}) #Renderizar vista pasando el formulario como contexto
    except:
        print("Las credenciales de usuario han expirado o existe algún problema con el ingreso")
        return render(request, "Usuarios/index.html")

def cveAcceso(request, token):
    decodedToken = jwt.decode(token, key=settings.SECRET_KEY, algorithms=['HS256'])
    if request.method=="POST":
        userp = User.objects.filter(username=request.POST['nomUsu'])[0].id
        print("userp ", userp)
        pacient = Paciente.objects.filter(user_id=userp)[0]
        print("pacient ", pacient)
        decodedToken = jwt.decode(token, key=settings.SECRET_KEY, algorithms=['HS256'])
        user = decodedToken['user_id']
        fecha = datetime.datetime.now()
        fechahoy= str(fecha.year)+"-"+str(fecha.month)+"-"+str(fecha.day)
        clave = str(fecha.day)+str(fecha.month)+str(fecha.year)[2:4]+str(fecha.hour)+str(fecha.minute)+str(user)+random.choice(string.ascii_uppercase)+str(random.randint(0,9))+random.choice(string.ascii_uppercase)
        print(clave)
        if Ap_Screening.objects.filter(resultadoFinal__isnull=True ,paciente=pacient): 
            print("No se pudo crear la sesión de tamizaje")
            return render(request, "Especialista/inicioEspecialista.html",{"exito": 'false', 'name': decodedToken['first_name'], 'access':token})
        else:
            screening = Ap_Screening.objects.create(cveAcceso=clave, paciente=pacient, fechaAp=fechahoy)
            screening.save()
            print(clave)
            return render(request, "Especialista/inicioEspecialista.html",{'clave':clave, 'name': decodedToken['first_name'], 'access':token })
    return render(request, "Especialista/inicioEspecialista.html", {'name': decodedToken['first_name'], 'access':token})

def reportes(request):
    user = '5'
    especialista = Especialista.objects.filter(user_id=user)[0].id
    #print(especialista)
    pacientes = []
    longitud = len(Paciente.objects.filter(especialista=especialista))
    #print(longitud)
    for i in range(0,longitud):        
        pacientes.append(Paciente.objects.filter(especialista=especialista)[i].id)
        #pruebas..append(pruebas)
    #print(pacientes)
    numero = len(Ap_Screening.objects.filter(paciente='1'))
    #print(numero)
    claves = []
    for i in pacientes:        
        clave = Ap_Screening.objects.filter(paciente=i)
        i = len(Ap_Screening.objects.filter(paciente=i))
        userP = Paciente.objects.filter(id=clave[0].paciente_id)[0].user_id
        nameP = User.objects.filter(id=userP)[0].first_name + " " + User.objects.filter(id=userP)[0].last_name
        for c in clave:
            claves.append({'index':i,'datos':c,'nombre':nameP})
            i=i-1
    pruebas = []
    return render(request, "Especialista/reportes.html", {'claves': claves})

def moca(request):
    if request.method=="POST":
        clave = request.POST.get('claveA')
        tamizaje = Screening.objects.filter(cveAcceso=clave)
        puntajes=[]
        imgs = []
        for i in tamizaje:
            puntaje = {'nReac':i.idReactivo, 'pts':i.puntajeReactivo}
            puntajes.append(puntaje)
            if i.idReactivo == 2:
                imagen = {'nReac':i.idReactivo, 'img': i.respuestaImg }
                imgs.append(imagen)
            elif i.idReactivo == 3:
                imagen = {'nReac':i.idReactivo, 'img': i.respuestaImg }
                imgs.append(imagen)
        
        print(puntajes)
        aplicacion = Ap_Screening.objects.filter(cveAcceso=clave)
        print(aplicacion[0].paciente_id)
        usuarioP = Paciente.objects.filter(id=aplicacion[0].paciente_id)[0].user_id
        nombreP = User.objects.filter(id=usuarioP)[0].first_name + " " + User.objects.filter(id=usuarioP)[0].last_name
        datos = {'fecha':str(aplicacion[0].fechaAp), 'nombre': nombreP}
        print(datos)
        paciente = Paciente.objects.filter(id=aplicacion[0].paciente_id)    
        doc = docx.Document("C:/Users/galil/Documents/GitHub/TT-II/TT/Zeitgeistv3/Especialista/moca.docx")
        nac = str(paciente[0].fechaNac)
        if paciente[0].sexo == 'F':
            sexo = "Femenino"
        else:
            sexo = "Masculino"
        if paciente[0].sexo == 'M':
            estudios = 'Ninguna'
        elif paciente[0].escolaridad == 'PR':
            estudios = 'Primaria'
        elif paciente[0].escolaridad == 'SC':
            estudios = 'Secundaria'
        elif paciente[0].escolaridad == 'BH':
            estudios = 'Bachillerato'
        else:
            estudios = 'Licenciatura o superior'

        fecha = str(aplicacion[0].fechaAp)
        print(fecha)
        #import win32.client
        #from win32.client import DispatchEx
        pythoncom.CoInitialize()#Plus the #
        #word = DispatchEx("Word.Application")
        #pythoncom.CoInitialize()#Plus the #
        
        p1 = doc.add_paragraph()
        p1.add_run('Nombre: ').bold = True
        p1.add_run(nombreP)
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2 = doc.add_paragraph()
        p2.add_run('Fecha de nac: ').bold = True
        p2.add_run(nac)
        p2.add_run('   Sexo: ').bold = True
        p2.add_run(sexo)
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3 = doc.add_paragraph()
        p3.add_run('Nivel de estudios: ').bold = True
        p3.add_run(estudios)
        p3.add_run('    Fecha: ').bold = True
        p3.add_run(fecha)
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph('			         '+'1'+'				          '+'1'+'				   '+'1'+'	   '+'5')
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph('			         '+'1'+'				          '+'1'+'				'+'1'+'	   '+'5')
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph('										         '+'1')
        doc.add_paragraph('										         '+'1'+'			    '+'2')
        doc.add_paragraph()
        doc.add_paragraph('						    '+'1'+'							   '+'1')
        doc.add_paragraph()
        doc.add_paragraph('						    							   '+'1')
        doc.add_paragraph('						  '+'1') # Lenguaje (parte 1)
        doc.add_paragraph('										'+'1'+'			   '+'2') 
        doc.add_paragraph('										 '+'1'+'			   '+'L') # Lenguaje (parte 2)
        p = doc.add_paragraph()
        p.size = Pt(10)
        pA = doc.add_paragraph('						          '+'1'+'		  '+'1'+'				   '+'2') #Abstracción
        pA.size = Pt(10)
        doc.add_paragraph('						    							   '+'5') # Recuerdo diferido (puntaje)
        p4 = doc.add_paragraph()
        p4.size = Pt(10)
        p4 = doc.add_paragraph()
        doc.add_paragraph('										         	     '+'1') # Recuerdo diferido. Palabras totales
        doc.add_paragraph()
        doc.add_paragraph('		 '+'*'+'	          '+'*'+'		     '+'*'+'		 '+'*'+'			'+'*'+'	       '+'*'+'		   '+'6') # Orientación
        doc.add_paragraph('									       '+'15') # MSI
        doc.add_paragraph('		'+'Galilea América Loretto Estrada') # Aplicador
        doc.add_paragraph('												         '+'26') # Total Final
        doc.save('C:/Users/galil/Documents/GitHub/TT-II/TT/Zeitgeistv3/staticfiles/reportes/moca-1.docx')
        #convert('moca-1.docx','moca.pdf')
        #pypandoc.convert_file('C:/Users/galil/Documents/GitHub/TT-II/TT/Zeitgeistv3/Especialista/reportes/moca-1.docx','pdf', outputfile="moca.pdf")
        #inFolder = 'C:/Users/galil/Documents/GitHub/TT-II/TT/Zeitgeistv3/Especialista/reportes/'
        #outFolder = 'C:/Users/galil/Documents/GitHub/TT-II/TT/Zeitgeistv3/Especialista/'
        #for fileName in os.listdir(inFolder):
            #print(fileName)
            #inFile = inFolder+fileName
            #word = comtypes.client.CreateObject('Word.Application')
            #word = win32com.client.DispatchEx("Word.Application")
            #doc = word.Documents.Open(inFile)
            #print('\n Abierto')
            #outFileName = inFile.replace('docx','pdf')
            #outFile = outFolder+outFileName
            #doc.SaveAs(outFile, FileFormat=17)
            #doc.Close()
            #word.Quit()
            #print('Se convirtió')"""
    return render(request, "Especialista/moca-pdf.html",{'datos':datos, 'imagenes':imgs})

def graphic(request):
    if request.method=="POST":
        clave = request.POST.get('clave')
        index = request.POST.get('index')
        paciente = Ap_Screening.objects.filter(cveAcceso=clave)[0].paciente_id
        userP = Paciente.objects.filter(id=paciente)[0].user_id
        nombreP = User.objects.filter(id=userP)[0].first_name + " " + User.objects.filter(id=userP)[0].last_name
        tamizajes = Ap_Screening.objects.filter(paciente=paciente)
        j = 1
        fechas = []
        for t in tamizajes:
            fechas.append(str(t.fechaAp))
        datos = {'nombre': nombreP, 'fecha1':fechas[1], 'fecha2':fechas[0]}
        tam = []
        i = len(tamizajes)
        if index == '2':
            for t in tamizajes:
                clave = i
                aplicaciones = Screening.objects.filter(cveAcceso=t.cveAcceso)
                for tamizaje in aplicaciones:
                    if tamizaje.idReactivo <= 3:
                        dato = {'cat':'vee','nReac':tamizaje.idReactivo, 'cal': tamizaje.puntajeReactivo } #visuoespacial/ejecutiva
                        tam.append({'cve': clave, 'datos': dato })                                
                    elif tamizaje.idReactivo == 4:
                        dato = {'cat':'ide','nReac':tamizaje.idReactivo, 'cal': tamizaje.puntajeReactivo } #identificacion
                        #datos.append(dato)
                        tam.append({'cve': clave, 'datos': dato })
                    elif tamizaje.idReactivo == 5:
                        dato = {'cat':'mem','nReac':tamizaje.idReactivo, 'cal': tamizaje.puntajeReactivo } #memoria
                        #datos.append(dato)
                        tam.append({'cve': clave, 'datos': dato })
                    elif 6 <= tamizaje.idReactivo <= 9 :
                        dato = {'cat':'atn', 'nReac':tamizaje.idReactivo, 'cal': tamizaje.puntajeReactivo } #atención
                        #datos.append(dato)
                        tam.append({'cve': clave, 'datos': dato })
                    elif 10 <= tamizaje.idReactivo <= 11:
                        dato = {'cat':'leng','nReac':tamizaje.idReactivo, 'cal': tamizaje.puntajeReactivo } #lenguaje
                        #datos.append(dato)
                        tam.append({'cve': clave, 'datos': dato })
                    elif tamizaje.idReactivo == 12 :
                        dato = {'cat':'abs', 'nReac':tamizaje.idReactivo, 'cal': tamizaje.puntajeReactivo } #abstración
                        #datos.append(dato)
                        tam.append({'cve': clave, 'datos': dato })
                    elif tamizaje.idReactivo == 13 :
                        dato = {'cat':'recd', 'nReac':tamizaje.idReactivo, 'cal': tamizaje.puntajeReactivo } #recuerdo diferido
                        #datos.append(dato)
                        tam.append({'cve': clave, 'datos': dato })
                    elif tamizaje.idReactivo == 14 :
                        dato = {'cat':'ort', 'nReac':tamizaje.idReactivo, 'cal': tamizaje.puntajeReactivo } #orientación
                        #datos.append(dato)
                        tam.append({'cve': clave, 'datos': dato })
                    
                i = i - 1
        #print(tam)
    return render(request, "Especialista/graficas.html",{'datos':datos,'mocas':tam})